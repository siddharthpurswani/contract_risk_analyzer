import re
import io
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

import PyPDF2
from docx import Document


# --- Output ---

@dataclass
class PreprocessedDocument:
    raw_text: str           # cleaned full text
    doc_type: str           # detected type: PARTNERSHIP, SERVICE_AGREEMENT, VENDOR, LICENSING, COMMERCIAL_CONTRACT, UNRECOGNIZED
    file_name: str
    total_lines: int
    warnings: list[str]     # e.g. OCR artifacts detected, low confidence type


# --- Doc type detection ---

CONTRACT_TYPE_KEYWORDS = {
    "PARTNERSHIP": ["partnership", "joint venture", "profit sharing", "partner", "collaboration agreement"],
    "SERVICE_AGREEMENT": ["service agreement", "scope of work", "deliverables", "service provider", "service fees"],
    "VENDOR": ["vendor", "supplier", "purchase order", "procurement", "goods and services"],
    "LICENSING": ["license", "licensee", "licensor", "royalty", "intellectual property license"],
}

# Clauses common to all commercial contracts — used to confirm it's a valid contract
COMMERCIAL_CLAUSE_KEYWORDS = [
    "liability", "termination", "confidentiality", "indemnity",
    "governing law", "dispute resolution", "payment", "warranty",
    "force majeure", "entire agreement", "amendment"
]

def detect_doc_type(text: str) -> str:
    text_lower = text.lower()

    # First check if it looks like a contract at all
    clause_hits = sum(1 for kw in COMMERCIAL_CLAUSE_KEYWORDS if kw in text_lower)
    if clause_hits < 2:
        return "UNRECOGNIZED"

    # Then identify the contract type
    scores = {ctype: 0 for ctype in CONTRACT_TYPE_KEYWORDS}
    for ctype, keywords in CONTRACT_TYPE_KEYWORDS.items():
        for kw in keywords:
            if kw in text_lower:
                scores[ctype] += 1

    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "COMMERCIAL_CONTRACT"


# --- Extractors ---

def extract_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def extract_from_pdf(file_bytes: bytes) -> str:
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def extract_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


# --- Cleaning ---

def clean_text(text: str) -> tuple[str, list[str]]:
    warnings = []

    # Detect likely OCR artifacts (unusual character clusters)
    ocr_pattern = r'[^\x00-\x7F]{3,}'
    if re.search(ocr_pattern, text):
        warnings.append("Possible OCR artifacts detected — non-ASCII clusters found.")

    # Remove page numbers (standalone numbers on a line)
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)

    # Remove common header/footer patterns
    text = re.sub(r'(Page \d+ of \d+)', '', text, flags=re.IGNORECASE)
    # Only strip CONFIDENTIAL/DRAFT when standalone on a line (headers/footers, not inside clauses)
    text = re.sub(r'^\s*(CONFIDENTIAL|DRAFT|FOR INTERNAL USE ONLY)\s*$', '', text, flags=re.IGNORECASE | re.MULTILINE)

    # Collapse excessive whitespace / blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)

    # Strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    text = "\n".join(lines)

    return text.strip(), warnings


# --- Main entry point ---

def preprocess(file_bytes: bytes, file_name: str) -> PreprocessedDocument:
    ext = Path(file_name).suffix.lower()

    if ext == ".pdf":
        raw = extract_from_pdf(file_bytes)
    elif ext == ".docx":
        raw = extract_from_docx(file_bytes)
    elif ext == ".txt":
        raw = extract_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    cleaned, warnings = clean_text(raw)
    doc_type = detect_doc_type(cleaned)
    total_lines = len(cleaned.splitlines())

    return PreprocessedDocument(
        raw_text=cleaned,
        doc_type=doc_type,
        file_name=file_name,
        total_lines=total_lines,
        warnings=warnings
    )